from __future__ import annotations

import io
from pathlib import Path
from typing import Any

try:
    import fitz
except ImportError:  # pragma: no cover
    fitz = None
try:
    import pytesseract
    from pytesseract import Output
except ImportError:  # pragma: no cover
    pytesseract = None
    Output = None
from docx import Document
from PIL import Image, ImageEnhance, ImageFilter, ImageOps


MOCK_TEXT = "[OCR unavailable or unsuccessful. Manual transcription review required.]"
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp"}


def extract_file(path: str | Path) -> dict[str, Any]:
    """Extract one file without allowing OCR or parser errors to escape."""
    source = Path(path)
    result: dict[str, Any] = {
        "file_name": source.name, "extracted_text": "", "ocr_method": "none",
        "confidence_score": 0.0, "page_count": 1, "status": "complete", "pages": [], "errors": [],
    }
    try:
        suffix = source.suffix.lower()
        if suffix == ".pdf":
            return _extract_pdf(source, result)
        if suffix in IMAGE_EXTENSIONS:
            with Image.open(source) as image:
                page = _ocr_image(image.convert("RGB"), 1)
            return _merge_pages(result, [page])
        if suffix == ".docx":
            document = Document(source)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
            for table in document.tables:
                text += "\n" + "\n".join(" | ".join(cell.text for cell in row.cells) for row in table.rows)
            result.update(extracted_text=text.strip() or MOCK_TEXT, ocr_method="direct_docx",
                          confidence_score=100.0 if text.strip() else 0.0,
                          status="complete" if text.strip() else "manual_review")
            return result
        if suffix == ".txt":
            raw = source.read_bytes()
            for encoding in ("utf-8-sig", "utf-8", "gb18030", "big5"):
                try:
                    text = raw.decode(encoding)
                    break
                except UnicodeDecodeError:
                    text = ""
            result.update(extracted_text=text or raw.decode("utf-8", errors="replace"),
                          ocr_method="direct_text", confidence_score=100.0)
            return result
        raise ValueError(f"Unsupported file type: {suffix}")
    except Exception as exc:
        result.update(extracted_text=MOCK_TEXT, ocr_method="mock_ocr", status="manual_review")
        result["errors"].append(str(exc))
        return result


def _extract_pdf(source: Path, base: dict[str, Any]) -> dict[str, Any]:
    if fitz is None:
        raise RuntimeError("PyMuPDF is not installed")
    pages: list[dict[str, Any]] = []
    with fitz.open(source) as document:
        base["page_count"] = document.page_count
        for number in range(document.page_count):
            try:
                page = document.load_page(number)
                direct = page.get_text("text").strip()
                if len(direct) >= 40:
                    pages.append({"page": number + 1, "text": direct, "method": "direct_pdf",
                                  "confidence": 100.0, "status": "complete", "errors": []})
                    continue
                pixmap = page.get_pixmap(matrix=fitz.Matrix(300 / 72, 300 / 72), alpha=False)
                image = Image.open(io.BytesIO(pixmap.tobytes("png"))).convert("RGB")
                pages.append(_ocr_image(image, number + 1))
            except Exception as exc:
                pages.append({"page": number + 1, "text": MOCK_TEXT, "method": "mock_ocr",
                              "confidence": 0.0, "status": "manual_review", "errors": [str(exc)]})
    return _merge_pages(base, pages)


def _ocr_image(image: Image.Image, page_number: int) -> dict[str, Any]:
    errors: list[str] = []
    if pytesseract is None:
        return {"page": page_number, "text": MOCK_TEXT, "method": "mock_ocr", "confidence": 0.0,
                "status": "manual_review", "errors": ["Tesseract Python binding is unavailable"]}
    gray = ImageOps.grayscale(image)
    variants = [image, ImageEnhance.Contrast(gray).enhance(1.8),
                ImageEnhance.Contrast(gray).enhance(2.3).filter(ImageFilter.SHARPEN)]
    best = {"text": "", "confidence": 0.0, "method": "tesseract"}
    for attempt, variant in enumerate(variants, 1):
        for language in ("chi_sim+eng", "eng"):
            try:
                data = pytesseract.image_to_data(variant, lang=language, config="--psm 6", output_type=Output.DICT,
                                                 timeout=120)
                words, scores = [], []
                for word, score in zip(data.get("text", []), data.get("conf", [])):
                    word = str(word).strip()
                    try:
                        numeric = float(score)
                    except (TypeError, ValueError):
                        numeric = -1
                    if word:
                        words.append(word)
                    if numeric >= 0:
                        scores.append(numeric)
                text = " ".join(words).strip()
                confidence = sum(scores) / len(scores) if scores else 0.0
                if len(text) > len(best["text"]) or confidence > best["confidence"]:
                    best = {"text": text, "confidence": confidence, "method": f"tesseract_{language}"}
                if text and confidence >= 85:
                    break
            except Exception as exc:
                errors.append(f"attempt {attempt}/{language}: {exc}")
        if best["text"] and best["confidence"] >= 85:
            break
    if not best["text"]:
        return {"page": page_number, "text": MOCK_TEXT, "method": "mock_ocr", "confidence": 0.0,
                "status": "manual_review", "errors": errors or ["OCR returned no text"]}
    status = "complete" if best["confidence"] >= 70 else "manual_review"
    return {"page": page_number, "text": best["text"], "method": best["method"],
            "confidence": round(best["confidence"], 2), "status": status, "errors": errors}


def _merge_pages(base: dict[str, Any], pages: list[dict[str, Any]]) -> dict[str, Any]:
    base["pages"] = pages
    base["page_count"] = len(pages)
    base["extracted_text"] = "\n\n".join(f"[Page {p['page']}]\n{p['text']}" for p in pages)
    base["confidence_score"] = round(sum(p["confidence"] for p in pages) / max(1, len(pages)), 2)
    methods = {p["method"] for p in pages}
    base["ocr_method"] = methods.pop() if len(methods) == 1 else "mixed"
    base["status"] = "manual_review" if any(p["status"] != "complete" for p in pages) else "complete"
    base["errors"] = [error for page in pages for error in page.get("errors", [])]
    return base
