import subprocess
import tempfile
from pathlib import Path

import fitz
from docx import Document


ALLOWED_EXTENSIONS = {".txt", ".pdf", ".doc", ".docx"}
MAX_EXTRACTED_CHARACTERS = 120_000


class DocumentError(ValueError):
    pass


def extract_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise DocumentError("Supported formats are TXT, PDF, DOC, and DOCX.")

    try:
        if suffix == ".txt":
            text = _decode_text(content)
        elif suffix == ".pdf":
            text = _extract_pdf(content)
        elif suffix == ".docx":
            text = _extract_docx(content)
        else:
            text = _extract_legacy_doc(content)
    except DocumentError:
        raise
    except Exception as exc:
        raise DocumentError(f"Could not read {filename}: {exc}") from exc

    text = text.strip()
    if not text:
        raise DocumentError(
            f"No selectable text was found in {filename}. Scanned PDFs require OCR before upload."
        )
    if len(text) > MAX_EXTRACTED_CHARACTERS:
        raise DocumentError(
            f"Extracted text is too long ({len(text):,} characters); the limit is "
            f"{MAX_EXTRACTED_CHARACTERS:,}."
        )
    return text


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "gb18030", "big5"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentError("TXT encoding is unsupported; save the file as UTF-8.")


def _extract_pdf(content: bytes) -> str:
    with fitz.open(stream=content, filetype="pdf") as document:
        return "\n\n".join(page.get_text("text") for page in document)


def _extract_docx(content: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".docx") as handle:
        handle.write(content)
        handle.flush()
        document = Document(handle.name)
        blocks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                blocks.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(blocks)


def _extract_legacy_doc(content: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".doc") as handle:
        handle.write(content)
        handle.flush()
        try:
            result = subprocess.run(
                ["antiword", handle.name],
                capture_output=True,
                check=True,
                timeout=30,
            )
        except FileNotFoundError as exc:
            raise DocumentError("Legacy DOC support requires antiword on the server.") from exc
        except subprocess.CalledProcessError as exc:
            raise DocumentError("The legacy DOC file could not be parsed.") from exc
        return _decode_text(result.stdout)
